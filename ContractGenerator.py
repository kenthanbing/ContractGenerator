import docx,os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml.ns import qn
import json

doc = Document('ContractGenerator' + os.sep + 'letterhead.docx')

#设置run字体
def set_font(run):
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return

# 获取json文件数据
with open('ContractGenerator' + os.sep + 'contract_data.json', 'rb') as f:
    contract_data = json.load(f)

# 设置正文全局字体
doc.styles['Normal'].font.name = '微软雅黑'
doc.styles['Normal'].font.size = Pt(12)
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 写合同号
doc._body.clear_content()
contract_number = contract_data['contract_number']
contract_number_run = doc.add_paragraph(f'合同号：{contract_number}')
contract_number_run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# 写标题
doc.add_heading(level=0).add_run(contract_data['exhibition_title'])
doc.add_heading(level=0).add_run('合作协议')

# 写正文
doc.add_paragraph(f'甲方：{contract_data["party_a"]}')
doc.add_paragraph(f'乙方：{contract_data["party_b"]}')
doc.add_paragraph()

doc.add_paragraph('双方本着平等自愿，协商一致的原则，签订此协议。双方均应严格遵守。')
doc.add_paragraph()

doc.add_paragraph('第一条	展位销售代理')
doc.add_paragraph('    甲方同意将其在中国区代理的下列展览会销售代理权授予乙方：')
doc.add_paragraph(f'    名称：{contract_data["exhibition_title"]}')
doc.add_paragraph(f'    时间：{contract_data["exhibition_date"]}')
doc.add_paragraph(f'    地点：{contract_data["exhibition_venue"]}')
doc.add_paragraph()

doc.add_paragraph('第二条	乙方职责')
doc.add_paragraph('    1. 乙方利用自身的销售网络在全国积极拓展客户。乙方向甲方转送接收到的展位询价和订单。乙方应把甲方规定的注意事项和参展条款对展商解释。采取措施配合甲方获得展位订单。')
doc.add_paragraph('    2. 乙方不得以甲方名义从事给甲方带来不良影响的活动。所有展位报价和定单须经甲方确认。')
doc.add_paragraph('    3. 未经甲方书面授权，乙方不得做出对展览会的承诺、担保或保证、陈述。')
doc.add_paragraph()

doc.add_paragraph('第三条	甲方职责')
doc.add_paragraph('    1. 甲方应向乙方提供展位效果图、摊位配置、价目表、广告宣传册及其他有关展位推销的辅助资料。')
doc.add_paragraph('    2. 甲方应全力支持乙方进行展位的推销，甲方不主动向乙方客户推广该展会。')
doc.add_paragraph('    3. 除本协议另有规定外，如乙方客户直接向甲方询价或订购展位，甲方应将该客户转介乙方联系。')
doc.add_paragraph()

doc.add_paragraph('第四条	商标')
doc.add_paragraph('    甲方目前拥有和使用的商标、图案及其他标记，均属甲方产权，未经甲方特别以书面同意，乙方均不得直接或间接，全部或部分注册。即使甲方特别以书面同意乙方按某种方式使用，但在本协议期满或终止后，此种使用应随即停止并取消。关于上述权利，如发生任何争议或索赔，甲方有权立即单方面取消本协议并且不承担由此而产生的任何责任。')
doc.add_paragraph()

doc.add_paragraph('第五条	费用结算方式')
doc.add_paragraph('    1. 甲方给予乙方展位合作价格标准如下：')
doc.add_paragraph(f'        {contract_data["booth_fee"]}；')
doc.add_paragraph('        报名费、注册费、角摊费免收。')
doc.add_paragraph('    2. 甲方给予乙方人员合作价格标准如下：')
doc.add_paragraph(f'        {contract_data["travel_fee"]}。')
doc.add_paragraph(f'    3. 付款期限：乙方于本协议签订后于一周内向甲方支付定金\
人民币8000元/展位；于{contract_data["payment_due_date"]}前向甲方支付剩余全部展位费及人员用。')
doc.add_paragraph()

doc.add_paragraph('第六条	协议期限')
doc.add_paragraph(f'    本协议自签订之日起生效，有效期至{contract_data["contract_valid_until"]}。')
doc.add_paragraph()

doc.add_paragraph('第七条	保密')
doc.add_paragraph('    双方保证对从另一方取得且无法自公开渠道获得的商业秘密（技术信息、 经营信息及其他商业秘密）予以保密。未经该商业秘密的原提供方同意，一方不得向任何第三方泄露该商业秘密的全部或部分内容。 法律、法规另有规定或双方另有约定的除外。')
doc.add_paragraph()

doc.add_paragraph('第八条	违约责任')
doc.add_paragraph('    双方严格履行本协议项下各自的义务，如果双方中任何一方实质性违反本协议的任何条款，使得协议无法履行或履行已经没有意义或严重影响协议的履行进程的，另一方有权立即终止本协议项下的一切合作。')
doc.add_paragraph()

doc.add_paragraph('第九条	争议与仲裁')
doc.add_paragraph('    双方在履行本协议发生争议，须经友好协商解决。如果协商不一致，提交北京市仲裁委员会按法令规定的程序进行仲裁，仲裁裁决为终局裁决。仲裁费用由败诉方承担。')
doc.add_paragraph()

#写落款，可以设置每个单元格的宽，同列单元格宽度相同，如果定义了不同的宽度将以最大值准

table = doc.add_table(rows=6, cols=3)
table.cell(0, 0).width = table.cell(0, 2).width = Cm(40)

cell_0_0 = table.cell(0, 0)
cell_0_0.text = f'甲方：{contract_data["party_a"]}'

cell_0_0 = table.cell(0, 2)
cell_0_0.text = f'已方：{contract_data["party_b"]}'

cell_2_0 = table.cell(2, 0)
cell_2_2 = table.cell(2, 2)
cell_2_0.text = cell_2_2.text = '签字：'

cell_4_0 = table.cell(4, 0)
cell_4_2 = table.cell(4, 2)
cell_4_0.text = cell_4_2.text = '公章：'

cell_5_0 = table.cell(5, 0)
cell_5_2 = table.cell(5, 2)
cell_5_0.text = cell_5_2.text = f'日期：{contract_data["contract_date"]}'



doc.save('ContractGenerator' + os.sep + f'{contract_data["exhibition_title"]}合作协议.docx')